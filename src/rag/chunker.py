from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from src.config import CHUNK_SIZE, CHUNK_OVERLAP, DATA_RAW
import os

# ── FORMATOS SUPORTADOS ──────────────────────────────────
FORMATOS_SUPORTADOS = (".txt", ".pdf", ".docx", ".xlsx", ".csv")

def carregar_txt(caminho: str):
    with open(caminho, encoding="utf-8") as f:
        texto = f.read()
    return [Document(page_content=texto, metadata={"source": caminho})]

def carregar_pdf(caminho: str):
    from pypdf import PdfReader
    reader = PdfReader(caminho)
    texto = "\n\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
    return [Document(page_content=texto, metadata={"source": caminho})]

def carregar_docx(caminho: str):
    from docx import Document as DocxDocument
    doc = DocxDocument(caminho)
    texto = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=texto, metadata={"source": caminho})]

def carregar_xlsx(caminho: str):
    import os as _os
    ext = _os.path.splitext(caminho)[1].lower()
    if ext == '.csv':
        import csv
        documentos = []
        with open(caminho, encoding='utf-8') as f:
            reader = csv.DictReader(f)
            linhas = [" | ".join([f"{k}: {v}" for k, v in row.items() if v]) for row in reader]
        texto = "\n".join(linhas)
        if texto.strip():
            documentos.append(Document(page_content=texto, metadata={"source": caminho, "sheet": "csv"}))
        return documentos
    import openpyxl
    wb = openpyxl.load_workbook(caminho, data_only=True)
    documentos = []

    for nome_aba in wb.sheetnames:
        aba = wb[nome_aba]
        rows = list(aba.iter_rows(values_only=True))

        # Detecta se é roteiro de promotor
        # (tem padrão SEMANA + dias da semana)
        eh_roteiro = any(
            row[0] and "SEMANA" in str(row[0])
            for row in rows if row[0]
        )

        if eh_roteiro:
            # Processa como roteiro narrativo
            texto = processar_roteiro(rows, caminho, nome_aba)
        else:
            # Processa como planilha normal
            texto = processar_planilha_normal(rows, nome_aba)

        if texto.strip():
            texto_com_aba = f"[Roteiro: {nome_aba}]\n{texto}"
            documentos.append(Document(
                page_content=texto_com_aba,
                metadata={"source": caminho, "sheet": nome_aba}
            ))

    return documentos


def processar_roteiro(rows, caminho, nome_aba):
    """Converte roteiro tabular em texto narrativo para o RAG."""
    blocos = []
    i = 0

    while i < len(rows):
        row = rows[i]

        # Detecta linha de SEMANA
        if row[0] and "SEMANA" in str(row[0]):
            semana = str(row[0])

            # Pega as datas da linha da semana
            datas = []
            for v in row[1:]:
                if v:
                    if hasattr(v, 'strftime'):
                        datas.append(v.strftime("%d/%m/%Y"))
                    else:
                        datas.append(str(v))
                else:
                    datas.append("")

            # Pega os dias da próxima linha
            dias = []
            if i + 1 < len(rows):
                prox = rows[i + 1]
                for v in prox[1:]:
                    dias.append(str(v) if v else "")

            # Monta dicionário dia → lojas
            visitas = {d: [] for d in dias if d}

            # Lê as lojas
            j = i + 2
            while j < len(rows):
                r = rows[j]
                # Para quando encontrar próxima semana ou fim
                if r[0] and ("SEMANA" in str(r[0]) or
                             "ITINERÁRIO" in str(r[0])):
                    break
                # Lê cada dia
                for col_idx, dia in enumerate(dias):
                    if dia and col_idx + 1 < len(r):
                        loja = r[col_idx + 1]
                        if loja:
                            visitas[dia].append(str(loja))
                j += 1

            # Gera texto narrativo
            bloco = f"{semana}\n"
            for k, dia in enumerate(dias):
                if dia and visitas.get(dia):
                    data = datas[k] if k < len(datas) else ""
                    bloco += f"\n{semana} — {dia}"
                    if data:
                        bloco += f" ({data})"
                    bloco += ":\n"
                    for seq, loja in enumerate(visitas[dia], 1):
                        bloco += f"  {seq}ª visita: {loja}\n"

            blocos.append(bloco)
            i = j
        else:
            i += 1

    return "\n\n".join(blocos)


def processar_planilha_normal(rows, nome_aba):
    """Converte planilha comum em texto tabular."""
    linhas = []
    cabecalho = None
    for i, row in enumerate(rows):
        row_limpa = [str(c) if c is not None else "" for c in row]
        if any(c for c in row_limpa):
            if cabecalho is None and i == 0:
                cabecalho = " | ".join(row_limpa)
                linhas.append(f"Planilha: {nome_aba}")
                linhas.append(f"COLUNAS: {cabecalho}")
            else:
                linhas.append(" | ".join(row_limpa))
    return "\n".join(linhas)
def processar_pasta(pasta: str = None) -> list:
    if pasta is None:
        pasta = DATA_RAW
    documentos = []
    for nome in os.listdir(pasta):
        caminho = os.path.join(pasta, nome)
        ext = os.path.splitext(nome)[1].lower()
        if ext not in FORMATOS_SUPORTADOS:
            continue
        print(f"[CHUNKER] Carregando: {nome}")
        if ext == ".txt":
            docs = carregar_txt(caminho)
        elif ext == ".pdf":
            docs = carregar_pdf(caminho)
        elif ext == ".docx":
            docs = carregar_docx(caminho)
        elif ext in (".xlsx", ".csv"):
            docs = carregar_xlsx(caminho)
        else:
            continue
        documentos.extend(docs)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )
    chunks = splitter.split_documents(documentos)
    print(f"[CHUNKER] Total de chunks: {len(chunks)}")
    return chunks